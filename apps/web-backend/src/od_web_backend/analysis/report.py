"""报告渲染:企业风格 HTML → PDF(WeasyPrint 优先,缺系统库回退 xhtml2pdf)+ DOCX(python-docx)。

图片以 base64 内联进 HTML,PDF 两种引擎都能出图;DOCX 直接按文件路径插图。
"""
from __future__ import annotations
import base64
import logging
from io import BytesIO
from pathlib import Path

from django.conf import settings

from .prompts import cn

logger = logging.getLogger(__name__)


def _abs(rel: str) -> Path:
    return Path(settings.MEDIA_ROOT) / rel


def _data_uri(rel: str) -> str:
    try:
        p = _abs(rel)
        b = p.read_bytes()
        return 'data:image/jpeg;base64,' + base64.b64encode(b).decode('ascii')
    except Exception:
        return ''


def _esc(s: str) -> str:
    return (str(s).replace('&', '&amp;').replace('<', '&lt;')
            .replace('>', '&gt;').replace('"', '&quot;'))


def _fmt_dt(iso: str) -> str:
    return (iso or '').replace('T', ' ')[:16]


# --------------------------------------------------------------------------- #
# HTML(表格布局,兼容 xhtml2pdf 与 WeasyPrint)
# --------------------------------------------------------------------------- #
def build_report_html(job, analysis: dict) -> str:
    summary = job.summary or {}
    per = summary.get('per_class', {})
    entries = sorted(per.items(), key=lambda kv: kv[1], reverse=True)
    images = list(job.images.all().prefetch_related('detections'))

    img_blocks = []
    for i, im in enumerate(images):
        dets = list(im.detections.all())
        rows = ''.join(
            f'<tr><td>{cn(d.label)}</td><td>{d.label}</td>'
            f'<td class="mono">{d.confidence * 100:.1f}%</td>'
            f'<td class="mono">[{", ".join(map(str, d.bbox))}]</td></tr>'
            for d in dets
        )
        det_table = (
            f'<table class="det"><thead><tr><th>类别</th><th>标签</th><th>置信度</th>'
            f'<th>坐标 [x1,y1,x2,y2]</th></tr></thead><tbody>{rows}</tbody></table>'
            if dets else '<div class="no-det">未检出疑似病变区域</div>'
        )
        img_blocks.append(f'''
        <div class="img-block">
          <div class="img-title">影像 {i + 1} · {_esc(im.filename)}
            <span class="muted">检出 {len(dets)} 处 · {im.width}×{im.height}</span></div>
          <table class="pair"><tr>
            <td><img src="{_data_uri(im.original)}"/><div class="cap">原始影像</div></td>
            <td><img src="{_data_uri(im.result)}"/><div class="cap">检测标注</div></td>
          </tr></table>
          {det_table}
        </div>''')

    sections_html = ''.join(
        f'<div class="analysis"><h3>{_esc(s["title"])}</h3>'
        f'<p>{_esc(s["content"]).replace(chr(10), "<br/>")}</p></div>'
        for s in analysis.get('sections', [])
    )

    return f'''<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"/>
<style>
  @page {{ size: A4; margin: 1.6cm 1.4cm; }}
  body {{ font-family: cjk, "Microsoft YaHei", "PingFang SC", sans-serif; color:#1f2733; font-size:12px; }}
  .mono {{ font-family: cjk, monospace; }}
  .muted {{ color:#909aa6; font-weight:normal; font-size:10px; }}
  .header {{ border-bottom:2px solid #0d8b8a; padding-bottom:8px; margin-bottom:12px; }}
  .header h1 {{ font-size:17px; margin:0 0 4px; color:#0d8b8a; }}
  .header .sub {{ color:#56606d; font-size:11px; }}
  .meta {{ font-size:10.5px; color:#56606d; margin-top:6px; }}
  .kpis {{ width:100%; margin:10px 0 14px; }}
  .kpis td {{ border:1px solid #e6e9ef; background:#fafbfc; padding:8px 10px; text-align:center; width:25%; }}
  .kpis .n {{ font-size:18px; font-weight:bold; color:#0d8b8a; }}
  .kpis .l {{ font-size:10px; color:#56606d; }}
  h2.sec {{ font-size:13px; color:#0a6e6d; border-left:3px solid #0d8b8a; padding-left:8px; margin:16px 0 8px; }}
  .analysis {{ margin-bottom:8px; }}
  .analysis h3 {{ font-size:12px; margin:0 0 3px; color:#0a6e6d; }}
  .analysis p {{ margin:0; line-height:1.6; color:#333; }}
  .img-block {{ margin-bottom:14px; }}
  .img-title {{ font-weight:bold; font-size:12px; margin-bottom:5px; }}
  table.pair {{ width:100%; }}
  table.pair td {{ width:50%; text-align:center; vertical-align:top; padding:2px; }}
  table.pair img {{ width:100%; border:1px solid #24304a; }}
  .cap {{ font-size:10px; color:#56606d; margin-top:2px; }}
  table.det {{ width:100%; border-collapse:collapse; margin-top:5px; }}
  table.det th, table.det td {{ border:1px solid #e6e9ef; padding:4px 6px; font-size:10.5px; text-align:left; }}
  table.det th {{ background:#f4f6f9; }}
  .no-det {{ color:#909aa6; font-size:11px; padding:4px 0; }}
  .disclaimer {{ margin-top:14px; padding:8px 10px; background:#fff6ee; border:1px solid #f2d5b8; font-size:10.5px; color:#8a5a2b; }}
  .footer {{ margin-top:14px; border-top:1px solid #e6e9ef; padding-top:6px; font-size:9.5px; color:#909aa6; text-align:center; }}
</style></head><body>
<div class="header">
  <h1>脑部 MRI 肿瘤检测报告</h1>
  <div class="sub">AI 辅助影像检测 · 仅供临床参考</div>
  <div class="meta">报告编号:{job.id} &nbsp;|&nbsp; 生成时间:{_fmt_dt(analysis.get("created_at", ""))}
    &nbsp;|&nbsp; 检测模型:{_esc(job.model)} &nbsp;|&nbsp; 分析引擎:{_esc(analysis.get("llm_model", ""))}</div>
</div>

<table class="kpis"><tr>
  <td><div class="n">{job.image_count}</div><div class="l">影像数量</div></td>
  <td><div class="n">{summary.get("count", 0)}</div><div class="l">检出目标总数</div></td>
  <td><div class="n">{len(entries)}</div><div class="l">涉及类别</div></td>
  <td><div class="n">{cn(entries[0][0]) if entries else "—"}</div><div class="l">主要类别</div></td>
</tr></table>

<h2 class="sec">一、智能分析</h2>
{sections_html}

<h2 class="sec">二、逐例影像检测</h2>
{''.join(img_blocks)}

<div class="disclaimer">⚠️ 免责声明:本报告由 AI 辅助检测系统自动生成,检测与分析结果仅供临床参考,<b>不能替代专业医师的诊断</b>。任何诊疗决策应由具备资质的医师结合完整临床资料作出。</div>
<div class="footer">ODPlatform 脑部 MRI 肿瘤检测系统 · 本页由系统自动生成</div>
</body></html>'''


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def _find_cjk_font() -> str | None:
    if settings.VIZ_FONT_PATH and Path(settings.VIZ_FONT_PATH).exists():
        return settings.VIZ_FONT_PATH
    for c in [
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Black.ttc',
        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
        '/usr/share/fonts/truetype/arphic/ukai.ttc',
    ]:
        if Path(c).exists():
            return c
    return None


def _register_cjk_xhtml2pdf() -> None:
    """给 xhtml2pdf(reportlab)注册 CJK 字体,让中文能渲染。"""
    font_path = _find_cjk_font()
    if not font_path:
        return
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        if 'cjk' not in pdfmetrics.getRegisteredFontNames():
            try:
                pdfmetrics.registerFont(TTFont('cjk', font_path))
            except Exception:
                pdfmetrics.registerFont(TTFont('cjk', font_path, subfontIndex=0))
    except Exception as exc:
        logger.warning('CJK 字体注册失败(xhtml2pdf): %s', exc)


def render_pdf(html: str) -> bytes:
    engine = settings.REPORT_PDF_ENGINE
    if engine in ('auto', 'weasyprint'):
        try:
            from weasyprint import HTML  # 需系统库(cairo/pango)
            return HTML(string=html).write_pdf()
        except Exception as exc:
            if engine == 'weasyprint':
                raise
            logger.info('WeasyPrint 不可用,回退 xhtml2pdf: %s', exc)

    from xhtml2pdf import pisa
    _register_cjk_xhtml2pdf()
    out = BytesIO()
    pisa.CreatePDF(src=html, dest=out, encoding='utf-8')
    return out.getvalue()


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def render_docx(job, analysis: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    summary = job.summary or {}
    per = summary.get('per_class', {})
    entries = sorted(per.items(), key=lambda kv: kv[1], reverse=True)
    images = list(job.images.all().prefetch_related('detections'))

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)

    title = doc.add_heading('脑部 MRI 肿瘤检测报告', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x8B, 0x8A)
    doc.add_paragraph('AI 辅助影像检测 · 仅供临床参考')

    meta = doc.add_paragraph()
    meta.add_run(
        f'报告编号:{job.id}\n生成时间:{_fmt_dt(analysis.get("created_at", ""))}\n'
        f'检测模型:{job.model}\n分析引擎:{analysis.get("llm_model", "")}'
    ).font.size = Pt(9)

    # KPI 表
    kt = doc.add_table(rows=2, cols=4)
    kt.style = 'Light Grid Accent 1'
    headers = ['影像数量', '检出目标总数', '涉及类别', '主要类别']
    values = [str(job.image_count), str(summary.get('count', 0)), str(len(entries)),
              cn(entries[0][0]) if entries else '—']
    for j, (h, v) in enumerate(zip(headers, values)):
        kt.cell(0, j).text = h
        kt.cell(1, j).text = v

    # 分析
    doc.add_heading('一、智能分析', level=1)
    for s in analysis.get('sections', []):
        doc.add_heading(s['title'], level=2)
        doc.add_paragraph(s['content'])

    # 逐例影像
    doc.add_heading('二、逐例影像检测', level=1)
    for i, im in enumerate(images):
        dets = list(im.detections.all())
        doc.add_heading(f'影像 {i + 1} · {im.filename}(检出 {len(dets)} 处)', level=2)
        # 原图 + 标注图
        pt = doc.add_table(rows=1, cols=2)
        for cell, rel, cap in [(pt.cell(0, 0), im.original, '原始影像'),
                               (pt.cell(0, 1), im.result, '检测标注')]:
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = para.add_run()
                run.add_picture(str(_abs(rel)), width=Inches(2.7))
            except Exception:
                para.add_run('(图片缺失)')
            cap_p = cell.add_paragraph(cap)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.runs[0].font.size = Pt(8)
        if dets:
            dt = doc.add_table(rows=1, cols=4)
            dt.style = 'Light List Accent 1'
            for j, h in enumerate(['类别', '标签', '置信度', '坐标']):
                dt.cell(0, j).text = h
            for d in dets:
                row = dt.add_row().cells
                row[0].text = cn(d.label)
                row[1].text = d.label
                row[2].text = f'{d.confidence * 100:.1f}%'
                row[3].text = f'[{", ".join(map(str, d.bbox))}]'
        else:
            doc.add_paragraph('未检出疑似病变区域')

    doc.add_paragraph()
    disc = doc.add_paragraph()
    r = disc.add_run(
        '⚠️ 免责声明:本报告由 AI 辅助检测系统自动生成,检测与分析结果仅供临床参考,'
        '不能替代专业医师的诊断。任何诊疗决策应由具备资质的医师结合完整临床资料作出。'
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x8A, 0x5A, 0x2B)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
